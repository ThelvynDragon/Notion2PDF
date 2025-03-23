import os
import markdown2
from bs4 import BeautifulSoup
from weasyprint import HTML
from docx import Document

def convert_md_to_html(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        markdown_text = f.read()
    html = markdown2.markdown(markdown_text, extras=["metadata", "fenced-code-blocks"])
    return html

def fix_internal_links(html, all_pages):
    soup = BeautifulSoup(html, "html.parser")
    for link in soup.find_all('a'):
        href = link.get('href', '')
        if href.endswith(".md"):
            page_name = os.path.splitext(os.path.basename(href))[0]
            if page_name in all_pages:
                link['href'] = f"#{page_name}"
    return str(soup)

def convert_zip_to_pdf_and_word(input_folder, selected_files, title, author, date, work_dir):
    all_pages = {os.path.splitext(f)[0]: f for f in selected_files}

    # PAGE DE COUVERTURE + SOMMAIRE
    full_html = f"""<html><head><meta charset='utf-8'>
    <style>{open("style.css", "r").read()}</style></head><body>
    <h1 style='text-align:center'>{title}</h1>
    <p style='text-align:center'><strong>{author}</strong> – {date}</p>
    <hr><h2>Table des matières</h2><ul>"""

    for page_name in all_pages:
        full_html += f"<li><a href='#{page_name}'>{page_name.replace('-', ' ')}</a></li>"
    full_html += "</ul><hr>"

    # DOCX INIT
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"{author} – {date}")
    doc.add_paragraph("------------------------------")

    for md_file in selected_files:
        page_name = os.path.splitext(md_file)[0]
        md_path = os.path.join(input_folder, md_file)
        html = convert_md_to_html(md_path)
        html_fixed = fix_internal_links(html, all_pages)

        full_html += f"<h1 id='{page_name}'>{page_name.replace('-', ' ')}</h1>" + html_fixed

        # Pour Word
        doc.add_heading(page_name.replace('-', ' '), level=1)
        with open(md_path, "r", encoding="utf-8") as f:
            doc.add_paragraph(f.read())

    full_html += "</body></html>"

    pdf_output = os.path.join(work_dir, "Notion_Document.pdf")
    docx_output = os.path.join(work_dir, "Notion_Document.docx")

    HTML(string=full_html, base_url=input_folder).write_pdf(pdf_output)
    doc.save(docx_output)

    return pdf_output, docx_output